from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import os
import shutil
import uuid
from pathlib import Path
from typing import Optional
import asyncio

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# ==================== FastAPI App Setup ====================
app = FastAPI(
    title="Luxury Law Firm - Document Engine",
    description="Local Python backend for document processing",
    version="1.0.0"
)

# CORS Configuration - Allow only local Tauri frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://127.0.0.1:*",
        "http://localhost:*",
        "tauri://localhost",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== Directory Setup ====================
BASE_DIR = Path.home() / ".luxury_law_editor"
ASSETS_DIR = BASE_DIR / "assets"
DOCUMENTS_DIR = BASE_DIR / "documents"
TEMP_DIR = BASE_DIR / "temp"

for directory in [ASSETS_DIR, DOCUMENTS_DIR, TEMP_DIR]:
    directory.mkdir(parents=True, exist_ok=True)

# Mount static files
app.mount("/assets", StaticFiles(directory=str(ASSETS_DIR)), name="assets")

# ==================== Data Models ====================
class MediaUploadResponse(BaseModel):
    url: str
    local_path: str
    file_name: str
    file_type: str
    size: int

class DocumentExportRequest(BaseModel):
    html_content: str
    output_format: str  # "docx", "pdf", "md", "epub", "html"
    document_title: str = "Document"
    author: str = "Author"
    file_name: Optional[str] = None

class DocumentConvertRequest(BaseModel):
    input_format: str
    output_format: str
    file_path: str

class SystemInfo(BaseModel):
    backend_version: str
    python_docx: bool
    pypandoc: bool
    pymupdf: bool
    assets_dir: str
    documents_dir: str

# ==================== Health & Info Endpoints ====================
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "Luxury Law Document Engine",
        "version": "1.0.0"
    }

@app.get("/api/system/info", response_model=SystemInfo)
async def get_system_info():
    """Get system information and available libraries"""
    return SystemInfo(
        backend_version="1.0.0",
        python_docx=PYTHON_DOCX_AVAILABLE,
        pypandoc=PANDOC_AVAILABLE,
        pymupdf=PYMUPDF_AVAILABLE,
        assets_dir=str(ASSETS_DIR),
        documents_dir=str(DOCUMENTS_DIR)
    )

# ==================== Media Upload Endpoints ====================
@app.post("/api/media/upload", response_model=MediaUploadResponse)
async def upload_media(file: UploadFile = File(...)):
    """Upload media files (images, videos)"""
    try:
        # Validate file
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        # Get file extension
        file_extension = Path(file.filename).suffix.lower()
        allowed_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.mp4', '.webm', '.avi', '.mov'}
        
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"File type {file_extension} not allowed")
        
        # Generate unique filename
        unique_name = f"{uuid.uuid4()}{file_extension}"
        dest_path = ASSETS_DIR / unique_name
        
        # Save file
        content = await file.read()
        with open(dest_path, "wb") as buffer:
            buffer.write(content)
        
        file_size = len(content)
        file_type = "image" if file_extension in {'.jpg', '.jpeg', '.png', '.gif', '.webp'} else "video"
        
        return MediaUploadResponse(
            url=f"http://127.0.0.1:8000/assets/{unique_name}",
            local_path=str(dest_path),
            file_name=unique_name,
            file_type=file_type,
            size=file_size
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.delete("/api/media/{file_name}")
async def delete_media(file_name: str):
    """Delete a media file"""
    try:
        file_path = ASSETS_DIR / file_name
        
        # Security check - prevent directory traversal
        if not file_path.resolve().is_relative_to(ASSETS_DIR.resolve()):
            raise HTTPException(status_code=403, detail="Invalid file path")
        
        if file_path.exists():
            file_path.unlink()
            return {"status": "success", "message": f"File {file_name} deleted"}
        else:
            raise HTTPException(status_code=404, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")

# ==================== Document Export Endpoints ====================
@app.post("/api/document/export")
async def export_document(payload: DocumentExportRequest):
    """Export document to various formats"""
    try:
        if not payload.output_format:
            raise HTTPException(status_code=400, detail="Output format is required")
        
        supported_formats = ["docx", "pdf", "md", "epub", "html"]
        if payload.output_format not in supported_formats:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {payload.output_format}")
        
        # Generate output filename
        file_name = payload.file_name or f"document_{uuid.uuid4()}"
        output_file = DOCUMENTS_DIR / f"{file_name}.{payload.output_format}"
        
        if payload.output_format == "docx":
            return await export_to_docx(payload, output_file)
        elif payload.output_format == "pdf":
            return await export_to_pdf(payload, output_file)
        elif payload.output_format in ["md", "epub"]:
            return await export_with_pandoc(payload, output_file)
        else:
            # HTML export - just save as is
            output_file.write_text(payload.html_content, encoding="utf-8")
            return {
                "status": "success",
                "message": f"Document exported to HTML",
                "file_path": str(output_file),
                "file_name": output_file.name
            }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

async def export_to_docx(payload: DocumentExportRequest, output_file: Path):
    """Export HTML content to DOCX format"""
    if not PYTHON_DOCX_AVAILABLE:
        raise HTTPException(status_code=400, detail="python-docx not installed")
    
    try:
        doc = DocxDocument()
        doc.add_heading(payload.document_title, 0)
        
        # Simple HTML to DOCX conversion (basic)
        from html.parser import HTMLParser
        
        class HTMLToDocxParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.document = doc
                self.current_paragraph = None
            
            def handle_starttag(self, tag, attrs):
                if tag == 'p':
                    self.current_paragraph = self.document.add_paragraph()
                elif tag == 'h1':
                    self.document.add_heading('', level=1)
                elif tag == 'h2':
                    self.document.add_heading('', level=2)
                elif tag == 'img':
                    attrs_dict = dict(attrs)
                    if 'src' in attrs_dict:
                        try:
                            self.document.add_picture(attrs_dict['src'])
                        except:
                            pass
            
            def handle_data(self, data):
                if data.strip():
                    if self.current_paragraph is None:
                        self.current_paragraph = self.document.add_paragraph()
                    self.current_paragraph.add_run(data)
        
        parser = HTMLToDocxParser()
        parser.feed(payload.html_content)
        
        doc.save(str(output_file))
        
        return {
            "status": "success",
            "message": "Document exported to DOCX",
            "file_path": str(output_file),
            "file_name": output_file.name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")

async def export_to_pdf(payload: DocumentExportRequest, output_file: Path):
    """Export HTML content to PDF format"""
    # Using pypandoc with wkhtmltopdf or weasyprint as fallback
    if not PANDOC_AVAILABLE:
        raise HTTPException(status_code=400, detail="pypandoc not installed")
    
    try:
        # Try with pandoc
        pypandoc.convert_text(
            payload.html_content,
            to="pdf",
            format="html",
            outputfile=str(output_file)
        )
        
        return {
            "status": "success",
            "message": "Document exported to PDF",
            "file_path": str(output_file),
            "file_name": output_file.name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")

async def export_with_pandoc(payload: DocumentExportRequest, output_file: Path):
    """Export using Pandoc for MD, EPUB, etc."""
    if not PANDOC_AVAILABLE:
        raise HTTPException(status_code=400, detail="pypandoc not installed")
    
    try:
        format_map = {
            "md": "markdown",
            "epub": "epub"
        }
        
        output_format = format_map.get(payload.output_format, payload.output_format)
        
        pypandoc.convert_text(
            payload.html_content,
            to=output_format,
            format="html",
            outputfile=str(output_file)
        )
        
        return {
            "status": "success",
            "message": f"Document exported to {payload.output_format.upper()}",
            "file_path": str(output_file),
            "file_name": output_file.name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

# ==================== Document Conversion Endpoints ====================
@app.post("/api/document/convert")
async def convert_document(payload: DocumentConvertRequest):
    """Convert between document formats"""
    if not PANDOC_AVAILABLE:
        raise HTTPException(status_code=400, detail="pypandoc not installed")
    
    try:
        if not Path(payload.file_path).exists():
            raise HTTPException(status_code=404, detail="Source file not found")
        
        output_file = DOCUMENTS_DIR / f"converted_{uuid.uuid4()}.{payload.output_format}"
        
        pypandoc.convert_file(
            payload.file_path,
            payload.output_format,
            outputfile=str(output_file)
        )
        
        return {
            "status": "success",
            "message": f"Document converted to {payload.output_format}",
            "file_path": str(output_file),
            "file_name": output_file.name
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

# ==================== Document Read Endpoints ====================
@app.get("/api/document/read/{file_name}")
async def read_document(file_name: str):
    """Read document content"""
    try:
        file_path = DOCUMENTS_DIR / file_name
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found")
        
        if file_path.suffix == ".docx":
            if not PYTHON_DOCX_AVAILABLE:
                raise HTTPException(status_code=400, detail="python-docx not installed")
            
            doc = DocxDocument(str(file_path))
            content = "\n".join([p.text for p in doc.paragraphs])
            return {"content": content, "format": "docx"}
        
        elif file_path.suffix == ".pdf":
            if not PYMUPDF_AVAILABLE:
                raise HTTPException(status_code=400, detail="PyMuPDF not installed")
            
            doc = fitz.open(str(file_path))
            content = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                content += page.get_text()
            return {"content": content, "format": "pdf", "pages": len(doc)}
        
        else:
            # For text formats
            content = file_path.read_text(encoding="utf-8")
            return {"content": content, "format": file_path.suffix.strip(".")}
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Read failed: {str(e)}")

# ==================== File Management Endpoints ====================
@app.get("/api/documents/list")
async def list_documents():
    """List all saved documents"""
    try:
        documents = []
        for file_path in DOCUMENTS_DIR.glob("*"):
            if file_path.is_file():
                documents.append({
                    "file_name": file_path.name,
                    "file_type": file_path.suffix.strip("."),
                    "size": file_path.stat().st_size,
                    "created": file_path.stat().st_ctime,
                    "modified": file_path.stat().st_mtime
                })
        return {"status": "success", "documents": documents}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"List failed: {str(e)}")

@app.delete("/api/documents/{file_name}")
async def delete_document(file_name: str):
    """Delete a saved document"""
    try:
        file_path = DOCUMENTS_DIR / file_name
        
        # Security check
        if not file_path.resolve().is_relative_to(DOCUMENTS_DIR.resolve()):
            raise HTTPException(status_code=403, detail="Invalid file path")
        
        if file_path.exists():
            file_path.unlink()
            return {"status": "success", "message": f"Document {file_name} deleted"}
        else:
            raise HTTPException(status_code=404, detail="File not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")

@app.get("/api/documents/download/{file_name}")
async def download_document(file_name: str):
    """Download a document file"""
    try:
        file_path = DOCUMENTS_DIR / file_name
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(
            path=str(file_path),
            filename=file_name,
            media_type="application/octet-stream"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000, log_level="info")
